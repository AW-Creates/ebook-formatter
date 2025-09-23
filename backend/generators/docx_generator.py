import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from .text_parser import parse_text, extract_title_and_chapters

def generate_docx(text: str, template_name: str, title: str = None, author: str = None) -> io.BytesIO:
    """
    Generate a Word document from text content with specified template styling.
    
    Args:
        text: Raw text content
        template_name: Name of the styling template to use
        title: Book title (optional, will be extracted from content if not provided)
        author: Author name (optional)
    
    Returns:
        io.BytesIO: Word document data
    """
    # Parse the text content
    content_blocks = parse_text(text)
    extracted_title, chapters = extract_title_and_chapters(content_blocks)
    
    # Use provided title or extracted title
    book_title = title if title and title != 'Untitled Book' else extracted_title
    book_author = author if author and author != 'Anonymous' else 'Anonymous'
    
    # Create new document
    doc = Document()
    
    # Configure document settings
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Apply template-specific styling
    apply_template_styles(doc, template_name)
    
    # Add title page
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(book_title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add author
    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_paragraph.add_run(f"by {book_author}")
    author_run.font.size = Pt(16)
    
    # Add page break after title page
    doc.add_page_break()
    
    # Add content
    for chapter in chapters:
        # Add chapter heading
        heading = doc.add_heading(chapter['title'], level=chapter['level'])
        apply_heading_style(heading, template_name, chapter['level'])
        
        # Add chapter content
        for paragraph_text in chapter['content']:
            paragraph = doc.add_paragraph(paragraph_text)
            apply_paragraph_style(paragraph, template_name)
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def apply_template_styles(doc: Document, template_name: str):
    """Apply template-specific document styles."""
    styles = doc.styles
    
    # Modify the Normal style based on template
    normal_style = styles['Normal']
    normal_font = normal_style.font
    
    if template_name == 'classic':
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        
    elif template_name == 'modern':
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
    elif template_name == 'elegant':
        normal_font.name = 'Georgia'
        normal_font.size = Pt(12)
        
    elif template_name == 'scifi':
        normal_font.name = 'Courier New'
        normal_font.size = Pt(10)
    
    # Set line spacing
    normal_style.paragraph_format.line_spacing = 1.15

def apply_heading_style(heading, template_name: str, level: int):
    """Apply template-specific heading styles."""
    heading_format = heading.paragraph_format
    heading_run = heading.runs[0] if heading.runs else None
    
    if not heading_run:
        return
    
    if template_name == 'classic':
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(18 - (level * 2))
        heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_format.space_before = Pt(24)
        heading_format.space_after = Pt(12)
        
    elif template_name == 'modern':
        heading_run.font.name = 'Calibri'
        heading_run.font.size = Pt(16 - (level * 1))
        heading_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_format.space_before = Pt(20)
        heading_format.space_after = Pt(10)
        
    elif template_name == 'elegant':
        heading_run.font.name = 'Georgia'
        heading_run.font.size = Pt(16 - (level * 1))
        heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_format.space_before = Pt(30)
        heading_format.space_after = Pt(15)
        heading_run.font.small_caps = True
        
    elif template_name == 'scifi':
        heading_run.font.name = 'Courier New'
        heading_run.font.size = Pt(14 - level)
        heading_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_format.space_before = Pt(16)
        heading_format.space_after = Pt(8)
        heading_run.font.bold = True
        heading_run.font.all_caps = True

def apply_paragraph_style(paragraph, template_name: str):
    """Apply template-specific paragraph styles."""
    paragraph_format = paragraph.paragraph_format
    
    if template_name == 'classic':
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.first_line_indent = Inches(0.5)
        paragraph_format.space_after = Pt(6)
        
    elif template_name == 'modern':
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_after = Pt(8)
        
    elif template_name == 'elegant':
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.first_line_indent = Inches(0.3)
        paragraph_format.space_after = Pt(10)
        
    elif template_name == 'scifi':
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_after = Pt(4)
        paragraph_format.left_indent = Inches(0.2)